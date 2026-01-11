from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Title
title = doc.add_heading('Learning Python - Complete Code Collection', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('A comprehensive collection of Python programs covering fundamentals and best practices.')
doc.add_paragraph('Organized in learning order - from basics to advanced concepts.')
doc.add_paragraph('By: Sai Mahir')
doc.add_paragraph()

# Table of Contents
doc.add_heading('Table of Contents', level=1)
toc = '''Section 1: Python Basics
    1.1 Syntax - First Python Program
    1.2 Keywords - Python Reserved Words
    1.3 Variable - Integer
    1.4 Variable - Float
    1.5 Variable - String
    1.6 Variable - Boolean
    1.7 Multiple Assignment
    1.8 Typecasting
    1.9 Format & Escape Characters

Section 2: Input and Output
    2.1 User Input
    2.2 Input Addition

Section 3: Strings
    3.1 String Methods
    3.2 String Indexing
    3.3 String Slicing

Section 4: Math Functions
    4.1 Math Module Functions

Section 5: Operators
    5.1 All Operators (Arithmetic, Relational, Logical, Bitwise)
    5.2 Membership & Identity Operators

Section 6: Conditional Statements
    6.1 Positive/Negative Check
    6.2 Greatest Integer - Using If-Elif
    6.3 Greatest Integer - Using max()
    6.4 Grade Calculator
    6.5 Grade Calculator with Average

Section 7: Loops
    7.1 Print N Numbers with Even/Odd Count
    7.2 Fibonacci Series
    7.3 Number Length
    7.4 Sum of Even Digits

Section 8: Practical Programs
    8.1 Swap Two Numbers
    8.2 Simple Interest
    8.3 Compound Interest
    8.4 Celsius to Fahrenheit
    8.5 Fahrenheit to Celsius
    8.6 Sum of First and Last Digit
    8.7 Calculator Program

Section 9: Number Theory Programs
    9.1 Armstrong Number
    9.2 Neon Number
    9.3 Harshad Number
    9.4 Palindrome - Using While Loop
    9.5 Palindrome - Using String Slicing

Section 10: Lists
    10.1 List Methods & Functions
    10.2 Arithmetic Operations on Lists

Section 11: Functions
    11.1 Basic Function - Addition
    11.2 Function with Numbers

Section 12: Problem Solving
    12.1 Count Element Occurrences
    12.2 Find Repeated Elements & Negatives
    12.3 Count with Validation
'''
doc.add_paragraph(toc)
doc.add_page_break()

# ============== SECTION 1: BASICS ==============
doc.add_heading('Section 1: Python Basics', level=1)

# 1.1 Syntax
doc.add_heading('1.1 Syntax (syntax.py)', level=2)
doc.add_paragraph('Summary: Your first Python program! The print() function outputs text to the console. Strings can be stored in variables.')
code = doc.add_paragraph()
code.add_run('''print("hello world")
print("i like pizza.")
s = "good morning"''').font.name = 'Consolas'

doc.add_paragraph()

# 1.2 Keywords
doc.add_heading('1.2 Keywords (keywords.py)', level=2)
doc.add_paragraph('Summary: Python has reserved words (keywords) that cannot be used as variable names. Use the keyword module to see all keywords.')
code = doc.add_paragraph()
code.add_run('''import keyword
print(keyword.kwlist)''').font.name = 'Consolas'

doc.add_paragraph()

# 1.3 Variable Integer
doc.add_heading('1.3 Variable - Integer (variableint.py)', level=2)
doc.add_paragraph('Summary: Integers are whole numbers. Use str() to concatenate numbers with strings. The type() function shows the data type.')
code = doc.add_paragraph()
code.add_run('''age = 18
age = 18 + 1
print(age) 
print("my age is "+ str(age))
print(type(age))''').font.name = 'Consolas'

doc.add_paragraph()

# 1.4 Variable Float
doc.add_heading('1.4 Variable - Float (variablefloat.py)', level=2)
doc.add_paragraph('Summary: Floats are decimal numbers. They require str() conversion for string concatenation.')
code = doc.add_paragraph()
code.add_run('''hight = 167.8
print(hight)
print("my hight is "+ str (hight))
print(type(hight))''').font.name = 'Consolas'

doc.add_paragraph()

# 1.5 Variable String
doc.add_heading('1.5 Variable - String (variablestring.py)', level=2)
doc.add_paragraph('Summary: Strings are text enclosed in quotes. Use + to concatenate (join) strings together.')
code = doc.add_paragraph()
code.add_run('''name = 'Sai'
lastname = "mahir"
fullname = name + lastname
print(fullname)
print("my name is "+name)
print(type(name))''').font.name = 'Consolas'

doc.add_paragraph()

# 1.6 Variable Boolean
doc.add_heading('1.6 Variable - Boolean (variableboolean.py)', level=2)
doc.add_paragraph('Summary: Booleans store True or False values. Useful for conditions and logical operations.')
code = doc.add_paragraph()
code.add_run('''donkey = True
print(donkey)
print("are you a donkey: "+ str(donkey))
print(type(donkey))''').font.name = 'Consolas'

doc.add_paragraph()

# 1.7 Multiple Assignment
doc.add_heading('1.7 Multiple Assignment (multipleAssignment.py)', level=2)
doc.add_paragraph('Summary: Python allows assigning multiple variables in one line, separated by commas. You can also assign the same value to multiple variables.')
code = doc.add_paragraph()
code.add_run('''name,age,hight="sai",18,167.8
print(name)
print(age)
print(hight)

sai = mahir = 18
print(sai)''').font.name = 'Consolas'

doc.add_paragraph()

# 1.8 Typecasting
doc.add_heading('1.8 Typecasting (typecasting.py)', level=2)
doc.add_paragraph('Summary: Typecasting converts one data type to another. Use int(), float(), str() functions. Essential for mixing different data types.')
code = doc.add_paragraph()
code.add_run('''a = 1
b = 2.5
c = "3"

a=float(a)   # int to float
b=str(b)     # float to string
c=int(c)     # string to int

print(a)
print(b)
print(c)

print("value of x is"+b)
# Typecasting is converting one data type to another.
# Required when concatenating strings with numbers.''').font.name = 'Consolas'

doc.add_paragraph()

# 1.9 Format & Escape Characters
doc.add_heading('1.9 Format & Escape Characters (format.py)', level=2)
doc.add_paragraph('Summary: Escape characters create special formatting. \\n = new line, \\t = tab space.')
code = doc.add_paragraph()
code.add_run('''print("my name is sai mahir\\ni am learning python")
print("my name is sai mahir\\t i am learning python")''').font.name = 'Consolas'

doc.add_paragraph()

# ============== SECTION 2: INPUT/OUTPUT ==============
doc.add_heading('Section 2: Input and Output', level=1)

# 2.1 Input
doc.add_heading('2.1 User Input (input.py)', level=2)
doc.add_paragraph('Summary: The input() function takes user input as a string. Use int() or float() to convert to numbers. Always convert back to str() for concatenation.')
code = doc.add_paragraph()
code.add_run('''name = input("Enter your name: ")
age = int(input("Enter your age: "))
height = float(input("Enter your height in meters: "))

age = str(age)
height = str(height)

print("hello "+name)
print("You are "+age+" years old.")
print("Your height is "+height+" meters.")''').font.name = 'Consolas'

doc.add_paragraph()

# 2.2 Input Addition
doc.add_heading('2.2 Input Addition (inputADD.py)', level=2)
doc.add_paragraph('Summary: Taking two numbers as input and performing addition. The int() conversion is crucial - without it, inputs are strings!')
code = doc.add_paragraph()
code.add_run('''a = int(input("enter first number: "))
b = int(input("enter second number: "))

sum = a+b

print("the sum is: ",sum)''').font.name = 'Consolas'

doc.add_paragraph()

# ============== SECTION 3: STRINGS ==============
doc.add_heading('Section 3: Strings', level=1)

# 3.1 String Basics
doc.add_heading('3.1 String Methods (stringbasics.py)', level=2)
doc.add_paragraph('Summary: Strings have many built-in methods. len() returns length, find() locates characters, upper()/lower() change case, isdigit()/isalpha() check content type.')
code = doc.add_paragraph()
code.add_run('''name = 'sai'
print(len(name))          # Length: 3
print(name[0])            # First character: s
print(name.find('i'))     # Index of 'i': 2
print(name.capitalize())  # First letter capital: Sai
print(name.upper())       # All uppercase: SAI
print(name.lower())       # All lowercase: sai
print(name.isdigit())     # Is it digits only? False
print(name.isalpha())     # Is it letters only? True
print(name.count('a'))    # Count occurrences of 'a': 1
print(name.replace('a','o'))  # Replace 'a' with 'o': soi
print(name * 5)           # Repeat 5 times: saisaisaisaisai''').font.name = 'Consolas'

doc.add_paragraph()

# 3.2 Indexing
doc.add_heading('3.2 String Indexing (indexing.py)', level=2)
doc.add_paragraph('Summary: Strings can be accessed by index (position). Index starts at 0. Use [::-1] to reverse a string. Negative indices count from the end.')
code = doc.add_paragraph()
code.add_run('''name = 'saimahir'

# Reverse a string using slicing
name = name[::-1]

print(name)  # Output: rihamiassss''').font.name = 'Consolas'

doc.add_paragraph()

# 3.3 Slicing
doc.add_heading('3.3 String Slicing (Slicing.py)', level=2)
doc.add_paragraph('Summary: Slicing extracts parts of a string. Syntax: string[start:end] or slice(start, end). Negative indices count from the end.')
code = doc.add_paragraph()
code.add_run('''website1 = 'https//google.com'

# Create a slice object (start at 7, end at -4)
slice_obj = slice(7,-4)

print(website1[slice_obj])  # Output: google

# Direct slicing examples:
# string[0:5]   - characters 0 to 4
# string[2:]    - from index 2 to end
# string[:5]    - from start to index 4
# string[-4:]   - last 4 characters''').font.name = 'Consolas'

doc.add_paragraph()

# ============== SECTION 4: MATH ==============
doc.add_heading('Section 4: Math Functions', level=1)

# 4.1 Math Functions
doc.add_heading('4.1 Math Module Functions (mathfunc.py)', level=2)
doc.add_paragraph('Summary: The math module provides advanced mathematical functions. Import it with "import math". Common functions include sqrt(), ceil(), floor(), factorial(), trigonometry, and more.')
code = doc.add_paragraph()
code.add_run('''import math

pi = 3.14

print(round(pi))           # Round to nearest: 3
print(math.ceil(pi))       # Round up: 4
print(math.floor(pi))      # Round down: 3
print(abs(-pi))            # Absolute value: 3.14
print(pow(pi,2))           # Power: pi^2
print(math.sqrt(16))       # Square root: 4.0
print(math.factorial(5))   # 5! = 120
print(max(1,2,3,4))        # Maximum: 4
print(min(1,2,3,4))        # Minimum: 1
print(math.gcd(12,15))     # Greatest common divisor: 3
print(math.lcm(12,15))     # Least common multiple: 60
print(math.sin(math.pi/2)) # Sine of 90 degrees
print(math.cos(0))         # Cosine of 0: 1.0
print(math.degrees(math.pi/2))  # Radians to degrees: 90
print(math.radians(90))    # Degrees to radians''').font.name = 'Consolas'

doc.add_paragraph()

# ============== SECTION 5: OPERATORS ==============
doc.add_heading('Section 5: Operators', level=1)

# 5.1 Operators
doc.add_heading('5.1 All Operators (operators.py)', level=2)
doc.add_paragraph('Summary: Python has several operator types:\n• Arithmetic: +, -, *, /, %, **, //\n• Relational: ==, !=, >, <, >=, <=\n• Logical: and, or, not\n• Assignment: =, +=, -=, *=, /=\n• Bitwise: &, |, ^, ~, <<, >>')
code = doc.add_paragraph()
code.add_run('''a = int(input("Enter first number: "))
b = int(input("Enter second number: "))

# Arithmetic Operators
print("\\nArithmetic Operators")
print("a + b =", a + b)   # Addition
print("a - b =", a - b)   # Subtraction
print("a * b =", a * b)   # Multiplication
print("a / b =", a / b)   # Division (float result)
print("a % b =", a % b)   # Modulus (remainder)
print("a ** b =", a ** b) # Exponentiation (power)
print("a // b =", a // b) # Floor division (integer result)

# Relational (Comparison) Operators - Return True/False
print("\\nRelational Operators")
print("a == b :", a == b)  # Equal to
print("a != b :", a != b)  # Not equal to
print("a > b  :", a > b)   # Greater than
print("a < b  :", a < b)   # Less than
print("a >= b :", a >= b)  # Greater than or equal
print("a <= b :", a <= b)  # Less than or equal

# Logical Operators - Combine conditions
print("\\nLogical Operators")
print("a > 0 and b > 0 :", a > 0 and b > 0)  # Both true
print("a > 0 or b > 0  :", a > 0 or b > 0)   # At least one true
print("not(a > 0)      :", not(a > 0))       # Opposite

# Assignment Operators - Modify and assign
print("\\nAssignment Operators")
c = a
c += b  # c = c + b
print("c += b ->", c)
c -= b  # c = c - b
print("c -= b ->", c)

# Bitwise Operators - Work on binary representation
print("\\nBitwise Operators")
print("a & b =", a & b)    # AND
print("a | b =", a | b)    # OR
print("a ^ b =", a ^ b)    # XOR
print("~a =", ~a)          # NOT (complement)
print("a << 1 =", a << 1)  # Left shift
print("a >> 1 =", a >> 1)  # Right shift''').font.name = 'Consolas'

doc.add_paragraph()

# 5.2 Membership & Identity Operators
doc.add_heading('5.2 Membership & Identity Operators (operators[extra].py)', level=2)
doc.add_paragraph('Summary: Membership operators (in, not in) check if value exists in sequence. Identity operators (is, is not) check if two variables point to same object in memory.')
code = doc.add_paragraph()
code.add_run('''list1 = [10, 20, 30, 40]
str1 = "python"

a = 20
b = 20
c = list1
d = list1.copy()  # Creates a new list with same values

# Membership Operators - Check existence
print("Membership Operators")
print("20 in list1       :", 20 in list1)       # True
print("50 not in list1   :", 50 not in list1)   # True
print("'p' in str1       :", 'p' in str1)       # True
print("'z' not in str1   :", 'z' not in str1)   # True

# Identity Operators - Check if same object
print("\\nIdentity Operators")
print("a is b            :", a is b)       # True (same value, same object)
print("a is not b        :", a is not b)   # False
print("c is d            :", c is d)       # False (different objects)
print("c is not d        :", c is not d)   # True''').font.name = 'Consolas'

doc.add_paragraph()

# ============== SECTION 6: CONDITIONALS ==============
doc.add_heading('Section 6: Conditional Statements', level=1)

# 6.1 Positive/Negative
doc.add_heading('6.1 Positive/Negative Check (positiveNegitive.py)', level=2)
doc.add_paragraph('Summary: if-elif-else statements execute code based on conditions. Conditions can be nested. Only one branch executes.')
code = doc.add_paragraph()
code.add_run('''a = int(input("Enter a number: "))

if a > 0:
    print("Positive")
    # Nested condition inside positive check
    if a % 2 == 0:
        print("Even")
    else:
        print("Odd")
elif a < 0:
    print("Negative")
    print("integer")
else:
    print("Zero")''').font.name = 'Consolas'

doc.add_paragraph()

# 6.2 Greatest Integer (if-elif)
doc.add_heading('6.2 Greatest Integer - Using If-Elif (greatestINT.py)', level=2)
doc.add_paragraph('Summary: Comparing multiple values using logical operators (and). Handles equal values with additional conditions.')
code = doc.add_paragraph()
code.add_run('''a = int(input("enter a number: "))
b = int(input("enter second number: "))
c = int(input("enter third number: "))

if a > b and a > c:
    print(a, "is the greatest integer")
elif b > a and b > c:
    print(b, "is the greatest integer")
elif a == b:
    print("First and second numbers are equal and greatest")
elif b == c:      
    print("Second and third numbers are equal and greatest")
elif a == b == c:
    print("All numbers are equal")
else:
    print(c, "is the greatest integer")''').font.name = 'Consolas'

doc.add_paragraph()

# 6.3 Greatest Integer (max)
doc.add_heading('6.3 Greatest Integer - Using max() (greatestinteger.py)', level=2)
doc.add_paragraph('Summary: The built-in max() function returns the largest value. Much simpler than multiple if-elif statements!')
code = doc.add_paragraph()
code.add_run('''a = int(input("Enter first number: "))
b = int(input("Enter second number: "))
c = int(input("Enter third number: "))

max_value = max(a, b, c)
print("Greatest integer is:", max_value)''').font.name = 'Consolas'

doc.add_paragraph()

# 6.4 Grades
doc.add_heading('6.4 Grade Calculator (Grades.py)', level=2)
doc.add_paragraph('Summary: Using range() with "in" operator to check if a value falls within a range. Clean way to handle multiple grade boundaries.')
code = doc.add_paragraph()
code.add_run('''a = int(input("enter your marks: "))

if a in range(90, 101):    # 90-100
    print("O")
elif a in range(80, 90):   # 80-89
    print("A+")
elif a in range(70, 80):   # 70-79
    print("A")
elif a in range(60, 70):   # 60-69
    print("B+") 
elif a in range(50, 60):   # 50-59
    print("B")
elif a in range(40, 50):   # 40-49
    print("C")
elif a in range(33, 40):   # 33-39
    print("D")  
else:                       # Below 33
    print("F")''').font.name = 'Consolas'

doc.add_paragraph()

# 6.5 Grades with Average
doc.add_heading('6.5 Grade Calculator with Average (grades1.py)', level=2)
doc.add_paragraph('Summary: Input validation (checking for invalid entries), calculating average of multiple values, and assigning grades based on the result.')
code = doc.add_paragraph()
code.add_run('''a = int(input("enter a number: "))
b = int(input("enter second number: "))
c = int(input("enter third number: "))
d = int(input("enter fourth number: "))
e = int(input("enter fifth number: "))

# Input validation - check for invalid marks
if a > 100 or b > 100 or c > 100 or d > 100 or e > 100:
    print("invalid entry")
elif a < 0 or b < 0 or c < 0 or d < 0 or e < 0:
    print("invalid entry")
else:
    # Calculate average
    sum = a + b + c + d + e
    avg = sum / 5

    # Assign grade based on average
    if avg >= 90:
        print("Grade O")
    elif avg >= 80:
        print("Grade A+")
    elif avg >= 70:
        print("Grade A")
    elif avg >= 60:
        print("Grade B+")
    elif avg >= 50:
        print("Grade B")
    elif avg >= 40:
        print("Grade C")
    elif avg >= 30:
        print("Grade D")
    else:
        print("Grade F")''').font.name = 'Consolas'

doc.add_paragraph()

# ============== SECTION 7: LOOPS ==============
doc.add_heading('Section 7: Loops', level=1)

# 7.1 N Numbers
doc.add_heading('7.1 Print N Numbers with Even/Odd Count (n-nUmber.py)', level=2)
doc.add_paragraph('Summary: for loop iterates over a sequence. range(1, n+1) generates numbers 1 to n. Use modulus (%) to check even/odd. Counters track occurrences.')
code = doc.add_paragraph()
code.add_run('''n = int(input("enter a number:"))
odd_count = 0
even_count = 0

for i in range(1, n+1):
    print(i)
    if i % 2 == 0:
        even_count = even_count + 1
    else:
        odd_count += 1  # Same as odd_count = odd_count + 1

print("even numbers=", even_count)
print("odd numbers=", odd_count)''').font.name = 'Consolas'

doc.add_paragraph()

# 7.2 Fibonacci
doc.add_heading('7.2 Fibonacci Series (fibonacci.py)', level=2)
doc.add_paragraph('Summary: Fibonacci sequence - each number is sum of previous two (0, 1, 1, 2, 3, 5, 8...). Uses variable swapping technique.')
code = doc.add_paragraph()
code.add_run('''# Fibonacci Series Generator
n = int(input("enter a number: "))

a = 0  # First number
b = 1  # Second number

for i in range(n):
    print(a)           # Print current number
    c = a + b          # Calculate next number
    a = b              # Move to next position
    b = c''').font.name = 'Consolas'

doc.add_paragraph()

# 7.3 Number Length
doc.add_heading('7.3 Number Length (numberLen.py)', level=2)
doc.add_paragraph('Summary: while loop continues until condition is false. Floor division (//) removes last digit. Count iterations to find number of digits.')
code = doc.add_paragraph()
code.add_run('''a = int(input("enter a number: "))
c = 0  # Counter for digits

# Using while loop
while a > 0:
    c = c + 1    # Increment counter
    a //= 10     # Remove last digit (same as a = a // 10)

print("length of number is:", c)

# Alternative: len(str(number)) gives same result''').font.name = 'Consolas'

doc.add_paragraph()

# 7.4 Sum of Even Digits
doc.add_heading('7.4 Sum of Even Digits (EvenSUm.py)', level=2)
doc.add_paragraph('Summary: Extract each digit using modulus (%), check if even, add to sum. Remove digit using floor division (//).')
code = doc.add_paragraph()
code.add_run('''n = int(input("enter a number:"))
even_sum = 0

while n > 0:
    digit = n % 10        # Get last digit
    if digit % 2 == 0:    # Check if digit is even
        even_sum = even_sum + digit
    n //= 10              # Remove last digit

print("sum of even digits is:", even_sum)''').font.name = 'Consolas'

doc.add_paragraph()

# ============== SECTION 8: PROGRAMS ==============
doc.add_heading('Section 8: Practical Programs', level=1)

# 8.1 Swap
doc.add_heading('8.1 Swap Two Numbers (swap.py)', level=2)
doc.add_paragraph('Summary: Swapping requires a temporary variable to hold one value. Python also allows: a, b = b, a')
code = doc.add_paragraph()
code.add_run('''a = int(input("enter first number: "))
b = int(input("enter second number: "))

# Method 1: Using temporary variable
temp = a    # Store a in temp
a = b       # Put b's value in a
b = temp    # Put original a (from temp) in b

print("after swapping the numbers are:", a, b)

# Method 2 (Python shortcut): a, b = b, a''').font.name = 'Consolas'

doc.add_paragraph()

# 8.2 Simple Interest
doc.add_heading('8.2 Simple Interest (simpleintrest.py)', level=2)
doc.add_paragraph('Summary: Simple Interest formula: SI = (P × R × T) / 100 where P=Principal, R=Rate, T=Time')
code = doc.add_paragraph()
code.add_run('''p = int(input("enter the principal amount: "))
r = float(input("enter the rate of interest: "))
t = int(input("enter the time period: "))

si = (p * r * t) / 100

print("the simple interest is:", si)''').font.name = 'Consolas'

doc.add_paragraph()

# 8.3 Compound Interest
doc.add_heading('8.3 Compound Interest (compundintrest.py)', level=2)
doc.add_paragraph('Summary: Compound Interest formula: A = P(1 + R/100)^T. CI = A - P (Final Amount minus Principal)')
code = doc.add_paragraph()
code.add_run('''P = float(input("Enter the principal amount: "))
R = float(input("Enter the rate of interest (%): "))
T = float(input("Enter the time (in years): "))

A = P * (1 + R / 100) ** T  # ** is exponentiation (power)

CI = A - P

print("Final Amount =", A)
print("Compound Interest =", CI)''').font.name = 'Consolas'

doc.add_paragraph()

# 8.4 Celsius to Fahrenheit
doc.add_heading('8.4 Celsius to Fahrenheit (~F-C.py)', level=2)
doc.add_paragraph('Summary: Temperature conversion formula: F = (9/5 × C) + 32')
code = doc.add_paragraph()
code.add_run('''c = int(input("enter the temperature in celsius:"))

f = (9/5 * c) + 32 

print("the temperature in fahrenheit is:", f)''').font.name = 'Consolas'

doc.add_paragraph()

# 8.5 Fahrenheit to Celsius
doc.add_heading('8.5 Fahrenheit to Celsius (~C-F.py)', level=2)
doc.add_paragraph('Summary: Temperature conversion formula: C = 5/9 × (F - 32)')
code = doc.add_paragraph()
code.add_run('''f = int(input("enter the temperature in fahrenheit: "))

c = (5/9) * (f - 32)

print("the temperature in celsius is:", c)''').font.name = 'Consolas'

doc.add_paragraph()

# 8.6 Sum of First and Last Digit
doc.add_heading('8.6 Sum of First and Last Digit (SUMf~l.py)', level=2)
doc.add_paragraph('Summary: Last digit = n % 10. First digit found by repeatedly dividing by 10 until single digit remains. Alternative: Convert to string and use indexing.')
code = doc.add_paragraph()
code.add_run('''n = int(input("enter a number:"))

# Method 1: Using while loop
last = n % 10      # Last digit using modulus
first = n 

while first >= 10:
    first //= 10   # Keep dividing until single digit

sum = first + last
print("Sum of first and last digit is:", sum)

# Method 2: Using string conversion
first_digit = int(str(n)[0])   # First character as int
last_digit = n % 10
sum_digits = first_digit + last_digit
print("Sum of first and last digit is:", sum_digits)''').font.name = 'Consolas'

doc.add_paragraph()

# 8.7 Calculator
doc.add_heading('8.7 Calculator Program (calc.py)', level=2)
doc.add_paragraph('Summary: A function-based calculator using if-elif to handle different operations. Functions are defined with def keyword.')
code = doc.add_paragraph()
code.add_run('''def calculator():
    sign = input("enter an operation (+, -, *, /): ")
    a = int(input("enter a number:"))
    b = int(input("enter a number:"))
    
    if sign == "+":
        result = a + b
        print(result)
    elif sign == "-":
        result = a - b
        print(result)
    elif sign == "*":
        result = a * b
        print(result)
    elif sign == "/":
        result = a / b
        print(result)
    else:
        print("wrong operator input")

calculator()  # Call the function''').font.name = 'Consolas'

doc.add_paragraph()

# ============== SECTION 9: NUMBER PROGRAMS ==============
doc.add_heading('Section 9: Number Theory Programs', level=1)

# 9.1 Armstrong
doc.add_heading('9.1 Armstrong Number (Armstrong.py)', level=2)
doc.add_paragraph('Summary: A number is Armstrong if sum of digits raised to power of number of digits equals the number itself.\nExample: 153 = 1³ + 5³ + 3³ = 1 + 125 + 27 = 153 ✓')
code = doc.add_paragraph()
code.add_run('''# Program to check Armstrong Number

num = int(input("Enter a number: "))
temp = num
sum = 0

# Find number of digits
digits = len(str(num))

# Calculate sum of digits raised to power
while temp > 0:
    digit = temp % 10
    sum += digit ** digits
    temp //= 10

# Check Armstrong condition
if sum == num:
    print(num, "is an Armstrong number")
else:
    print(num, "is not an Armstrong number")''').font.name = 'Consolas'

doc.add_paragraph()

# 9.2 Neon
doc.add_heading('9.2 Neon Number (neon.py)', level=2)
doc.add_paragraph('Summary: A number is Neon if sum of digits of its square equals the number.\nExample: 9² = 81, and 8 + 1 = 9 ✓')
code = doc.add_paragraph()
code.add_run('''a = int(input("enter a number:"))
n = a ** 2     # Square the number
sum = 0

# Sum the digits of the square
while n > 0:
    digit = n % 10
    sum = digit + sum
    n //= 10

if sum == a:
    print(a, "is a neon number")
else:
    print(a, "is not a neon number")''').font.name = 'Consolas'

doc.add_paragraph()

# 9.3 Harshad Number
doc.add_heading('9.3 Harshad Number (harshadnumber.py)', level=2)
doc.add_paragraph('Summary: A number is Harshad (Niven) if it is divisible by sum of its digits.\nExample: 18 → 1+8=9, 18÷9=2 (no remainder) ✓')
code = doc.add_paragraph()
code.add_run('''n = int(input("Enter a number: "))

sum = 0
temp = n

# Sum the digits
while n > 0:
    digit = n % 10
    sum = sum + digit
    n //= 10

# Check if divisible by digit sum
if temp % sum == 0:
    print("It is a Harshad number")
else:
    print("It is not a Harshad number")''').font.name = 'Consolas'

doc.add_paragraph()

# 9.4 Palindrome (Method 1)
doc.add_heading('9.4 Palindrome - Using While Loop (pallendrome.py)', level=2)
doc.add_paragraph('Summary: A palindrome reads the same forward and backward. Reverse the number by extracting digits and compare.\nExamples: 121, 12321, 1001')
code = doc.add_paragraph()
code.add_run('''a = int(input("enter a number: "))
rev = 0
temp = a

# Reverse the number
while temp > 0:
    digit = temp % 10           # Get last digit
    rev = rev * 10 + digit      # Build reversed number
    temp //= 10                 # Remove last digit

if rev == a:
    print(a, "is a palindrome number")
else:
    print(a, "is not a palindrome number")''').font.name = 'Consolas'

doc.add_paragraph()

# 9.5 Palindrome (Method 2)
doc.add_heading('9.5 Palindrome - Using String Slicing (pallendrome(1).py)', level=2)
doc.add_paragraph('Summary: Python strings can be reversed using [::-1] slice notation. Much simpler approach!')
code = doc.add_paragraph()
code.add_run('''a = str(input("enter a number: "))

temp = a[::-1]   # Reverse using slicing

if temp == a:
    print(a, "is a palindrome number")
else:
    print(a, "is not a palindrome number")''').font.name = 'Consolas'

doc.add_paragraph()

# ============== SECTION 10: LISTS ==============
doc.add_heading('Section 10: Lists', level=1)

# 10.1 List Functions
doc.add_heading('10.1 List Methods & Functions (listFunc.py)', level=2)
doc.add_paragraph('Summary: Lists are ordered, mutable collections. Key methods:\n• len() - length  • max()/min() - largest/smallest  • sum() - total\n• sort() - sort in place  • reverse() - reverse in place\n• append() - add item  • remove() - delete item  • pop() - remove & return last\n• insert() - add at index  • extend() - add multiple items  • count() - occurrences\nNote: Methods like remove(), sort(), clear() modify the list and return None!')
code = doc.add_paragraph()
code.add_run('''my_list = [1, 2, 3, 4, 5, 6, 7, 8, 9, 10]

print("Length:", len(my_list))           # 10
print("Max:", max(my_list))              # 10
print("Min:", min(my_list))              # 1
print("Sum:", sum(my_list))              # 55
print("Sorted (reverse):", sorted(my_list, reverse=True))

my_list.remove(5)                        # Remove value 5
print("After remove(5):", my_list)

print("Reversed:", list(reversed(my_list)))

my_list.insert(2, 15)                    # Insert 15 at index 2
print("After insert(2,15):", my_list)

print("Count of 3:", my_list.count(3))   # Count occurrences
print("Index of 4:", my_list.index(4))   # Find position

my_list.sort()                           # Sort ascending
print("After sort():", my_list)

popped = my_list.pop()                   # Remove last item
print("Popped element:", popped)

my_list.extend([11, 12, 13])             # Add multiple items
print("After extend:", my_list)

my_list.clear()                          # Remove all items
print("After clear():", my_list)''').font.name = 'Consolas'

doc.add_paragraph()

# 10.2 Arithmetic Lists
doc.add_heading('10.2 Arithmetic Operations on Lists (arthemeticList.py)', level=2)
doc.add_paragraph('Summary: Perform element-wise arithmetic operations on two lists. Uses map() to convert space-separated input into a list of integers.')
code = doc.add_paragraph()
code.add_run('''# Input lists using map() and split()
a = list(map(int, input("enter the elements: ").split()))
b = list(map(int, input("enter the elements: ").split()))
operator = input("Enter the operation (+, -, *, /): ")

# Create result lists
add = []
sub = []
multiply = []
division = []

# Perform operations element by element
for i in range(len(a)):
    add.append(a[i] + b[i])
    sub.append(a[i] - b[i])
    multiply.append(a[i] * b[i])
    division.append(a[i] / b[i])

# Display based on operator choice
if operator == '+':
    print("The addition of two lists is:", add)
elif operator == '-':
    print("The subtraction of two lists is:", sub)
elif operator == '*':
    print("The multiplication of two lists is:", multiply)
elif operator == '/':
    print("The division of two lists is:", division)
else:
    print("Invalid operator")''').font.name = 'Consolas'

doc.add_paragraph()

# ============== SECTION 11: FUNCTIONS ==============
doc.add_heading('Section 11: Functions', level=1)

# 11.1 Basic Function
doc.add_heading('11.1 Basic Function - Addition (addFunc.py)', level=2)
doc.add_paragraph('Summary: Functions are defined using "def" keyword. They group reusable code. Call the function by its name followed by parentheses.')
code = doc.add_paragraph()
code.add_run('''def add():
    a = int(input("enter a number:"))
    b = int(input("enter second number:"))
    print("Sum is:", a + b)

add()  # Call the function to execute it''').font.name = 'Consolas'

doc.add_paragraph()

# 11.2 Function with Numbers
doc.add_heading('11.2 Function with Numbers (funcNnum.py)', level=2)
doc.add_paragraph('Summary: Functions can take parameters. Here n is passed to the function, which then uses it in a loop.')
code = doc.add_paragraph()
code.add_run('''def print_numbers(n):
    for i in range(1, n+1):
        print(i)

n = int(input("enter a number:"))
print_numbers(n)  # Pass n as argument to the function''').font.name = 'Consolas'

doc.add_paragraph()

# ============== SECTION 12: PROBLEM SOLVING ==============
doc.add_heading('Section 12: Problem Solving', level=1)

# 12.1 Count Element
doc.add_heading('12.1 Count Element Occurrences (prob4(1).py)', level=2)
doc.add_paragraph('Summary: Count how many times an element appears in a list. Important: Always convert input to int when working with numbers!')
code = doc.add_paragraph()
code.add_run('''n = int(input("enter a value:"))

lst = []
count = 0

for i in range(n):
    lst.append(int(input("enter the element:")))  # Convert to int!

a = int(input("enter the number to be checked:"))

for i in lst:
    if i == a:
        count = count + 1

print("Element", a, "appears", count, "times")
# Alternative: print(lst.count(a))''').font.name = 'Consolas'

doc.add_paragraph()

# 12.2 Repeated Elements & Negatives
doc.add_heading('12.2 Find Repeated Elements & Negatives (prob4.py)', level=2)
doc.add_paragraph('Summary: Use a dictionary to count occurrences. Keys are elements, values are counts. Iterate through list to find negatives.')
code = doc.add_paragraph()
code.add_run('''n = int(input("Enter number of elements: "))
lst = []

for i in range(n):
    lst.append(int(input("Enter element: ")))

# Count occurrences using dictionary
count_dict = {}
neg_count = 0

for item in lst:
    if item in count_dict:
        count_dict[item] += 1
    else:
        count_dict[item] = 1

# Count negative elements
for i in range(n):
    if lst[i] < 0:
        neg_count = neg_count + 1

# Display repeated elements
print("\\nRepeated elements and their counts:")
for key, value in count_dict.items():
    if value > 1:
        print(key, "is repeated", value, "times")

print("\\nNumber of negative elements in the list:", neg_count)''').font.name = 'Consolas'

doc.add_paragraph()

# 12.3 Count with Validation
doc.add_heading('12.3 Count with Validation (listP3.py)', level=2)
doc.add_paragraph('Summary: Complete program to count occurrences and negatives with proper validation messages when elements are not found.')
code = doc.add_paragraph()
code.add_run("""'''WAP to count occurrences of an element in a list.
Also count number of negative elements in the list.
If element or negative elements are not found, print
appropriate message.'''

n = int(input("enter a value:"))

lst = []
count = 0
neg_count = 0

for i in range(n):
    lst.append(int(input("enter the element:")))

a = int(input("enter the number to be checked:"))

# Count occurrences of element
for i in lst:
    if i == a:
        count = count + 1

# Alternative: count = lst.count(a)

# Count negative elements
for i in lst:
    if i < 0:
        neg_count += 1

# Display results with validation
if count != 0:
    print("Element", a, "appears", count, "times")
else:
    print("No element found")

if neg_count != 0:
    print("Number of negative elements:", neg_count)
else:
    print("No negative element found")""").font.name = 'Consolas'

doc.add_paragraph()
doc.add_page_break()

# Final Notes
doc.add_heading('Quick Reference - Key Concepts', level=1)
notes = '''
DATA TYPES:
• int - whole numbers (1, 42, -5)
• float - decimal numbers (3.14, -0.5)
• str - text ("hello", 'world')
• bool - True or False
• list - ordered collection [1, 2, 3]

TYPE CONVERSION:
• int() - convert to integer
• float() - convert to float
• str() - convert to string
• list() - convert to list

INPUT/OUTPUT:
• input() - always returns string!
• print() - display output
• print(x, y) - multiple values
• print("text" + str(num)) - concatenation

OPERATORS:
• Arithmetic: +, -, *, /, %, **, //
• Comparison: ==, !=, >, <, >=, <=
• Logical: and, or, not
• Assignment: =, +=, -=, *=, /=

LOOPS:
• for i in range(n): - loop n times
• for item in list: - loop through list
• while condition: - loop while true

CONDITIONS:
• if condition:
• elif condition:
• else:

COMMON LIST METHODS:
• append(x) - add to end
• remove(x) - remove first occurrence
• pop() - remove & return last
• sort() - sort in place
• len(list) - get length
• list.count(x) - count occurrences

USEFUL PATTERNS:
• n % 10 - get last digit
• n // 10 - remove last digit
• str[::-1] - reverse string/number
• len(str(n)) - number of digits
'''
doc.add_paragraph(notes)

# Save the document
doc.save('C:/Users/saima/OneDrive/Documents/Sem 2/learning python/Python_Code_Collection.docx')
print("Document created successfully: Python_Code_Collection.docx")
