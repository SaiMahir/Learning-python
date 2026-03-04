from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Categorized files with descriptions
categories = {
    "1. Variables": [
        {"filename": "variableint.py", "description": "Demonstrates integer variable declaration, arithmetic operations, and type checking with print statements."},
        {"filename": "variablefloat.py", "description": "Shows float variable usage, printing values with string concatenation, and type() function."},
        {"filename": "variablestring.py", "description": "Demonstrates string variables, string concatenation to form full name, and type checking."},
        {"filename": "variableboolean.py", "description": "Shows boolean variable usage, printing boolean values with string concatenation."},
        {"filename": "multipleAssignment.py", "description": "Demonstrates multiple variable assignment in a single line and assigning same value to multiple variables."},
        {"filename": "keywords.py", "description": "Uses the keyword module to print all Python reserved keywords."},
        {"filename": "syntax.py", "description": "Basic Python syntax demonstration with print statements and string variable assignment."}
    ],
    "2. Input/Output": [
        {"filename": "input.py", "description": "Takes user input for name, age, and height, then displays them with formatted output."},
        {"filename": "inputADD.py", "description": "Takes two integer inputs from user and calculates their sum."},
        {"filename": "format.py", "description": "Demonstrates escape characters \\n (newline) and \\t (tab) in print statements."}
    ],
    "3. Operators": [
        {"filename": "operators.py", "description": "Comprehensive demonstration of arithmetic, relational, logical, assignment, and bitwise operators."},
        {"filename": "operators[extra].py", "description": "Demonstrates membership operators (in, not in) and identity operators (is, is not)."}
    ],
    "4. Type Casting": [
        {"filename": "typecasting.py", "description": "Shows type conversion between int, float, and string using float(), str(), and int() functions."}
    ],
    "5. Strings": [
        {"filename": "stringbasics.py", "description": "Demonstrates string methods: len, indexing, find, capitalize, upper, lower, isdigit, isalpha, count, replace, and repetition."},
        {"filename": "indexing.py", "description": "Reverses a string using slicing with [::-1] notation."},
        {"filename": "Slicing.py", "description": "Extracts a substring from a URL using the slice() function."}
    ],
    "6. Conditionals": [
        {"filename": "Grades.py", "description": "Assigns letter grades (O, A+, A, B+, B, C, D, F) based on marks using if-elif-else."},
        {"filename": "grades1.py", "description": "Calculates average of 5 subject marks, validates input, and assigns grades based on average."},
        {"filename": "positiveNegitive.py", "description": "Checks if a number is positive, negative, or zero, and whether it's even or odd."},
        {"filename": "greatestINT.py", "description": "Finds the greatest among three numbers using if-elif-else with equality checks."},
        {"filename": "greatestinteger.py", "description": "Finds the greatest of three numbers using max() function."}
    ],
    "7. Loops": [
        {"filename": "n-nUmber.py", "description": "Prints numbers from 1 to n and counts even and odd numbers."},
        {"filename": "swap.py", "description": "Swaps two numbers using a temporary variable."}
    ],
    "8. Functions": [
        {"filename": "addFunc.py", "description": "Defines a function to take two numbers as input and print their sum."},
        {"filename": "calc.py", "description": "Implements a calculator function with addition, subtraction, multiplication, and division operations."},
        {"filename": "evenoddFunc.py", "description": "Defines a function that returns True if a number is even, False otherwise."},
        {"filename": "funcNnum.py", "description": "Defines a function to print numbers from 1 to n using a loop."},
        {"filename": "lambda.py", "description": "Demonstrates lambda functions for square, add, subtract, multiply, divide, and finding greater number."},
        {"filename": "ReduceFilter.py", "description": "Shows filter() for filtering even numbers and reduce() for calculating product of list elements."}
    ],
    "9. Lists": [
        {"filename": "listFunc.py", "description": "Demonstrates list methods: len, max, min, sum, sorted, remove, reversed, insert, count, index, sort, pop, extend, clear."},
        {"filename": "listcopy.py", "description": "Shows how to append one list to another using append() method."},
        {"filename": "ListComprehention.py", "description": "Demonstrates list comprehension to convert Celsius temperatures to Fahrenheit."},
        {"filename": "dynamiclist.py", "description": "Creates a list dynamically from user input and calculates sum of even and odd numbers separately."},
        {"filename": "arthemeticList.py", "description": "Performs element-wise arithmetic operations (+, -, *, /) on two lists based on user-selected operator."},
        {"filename": "listP0.py", "description": "Filters a list to find prime, Armstrong, perfect, and neon numbers using defined functions."},
        {"filename": "listP1.py", "description": "Copies a list to another list using copy() method and other techniques."},
        {"filename": "listP2.py", "description": "Separates even and odd numbers from a list into two different lists."},
        {"filename": "listP3.py", "description": "Counts occurrences of an element and negative numbers in a list."},
        {"filename": "listP4.py", "description": "Creates a list and computes the square of each element."},
        {"filename": "listP5.py", "description": "Creates a dictionary from two lists where elements of first list are keys and second are values."},
        {"filename": "listP6.py", "description": "Counts frequency of each element in a list and stores in a dictionary."},
        {"filename": "listP7.py", "description": "Converts a list to a tuple."},
        {"filename": "listP8.py", "description": "Multiplies corresponding elements of two lists."},
        {"filename": "listP9.py", "description": "Removes duplicate elements from a list."},
        {"filename": "listP10.py", "description": "Finds common elements between two lists."},
        {"filename": "listP11.py", "description": "Filters out even numbers from a list, keeping only odd numbers."},
        {"filename": "listP12.py", "description": "Checks if each element in a list is prime or not."},
        {"filename": "listP13.py", "description": "Creates a list of lists by taking multiple lists as input."},
        {"filename": "listP14.py", "description": "Picks random elements from a list using random.choice()."},
        {"filename": "listP15.py", "description": "Sorts a list and finds the second smallest element."},
        {"filename": "listP16.py", "description": "Counts how many times each element is repeated in a list."},
        {"filename": "listP17.py", "description": "Checks if words from a list are present in a given sentence."},
        {"filename": "listP18.py", "description": "Replaces the last element of a list with a sublist."},
        {"filename": "listP19.py", "description": "Prints corresponding elements from two lists side by side."},
        {"filename": "listP20.py", "description": "Moves all zeros in a list to the end while maintaining order of other elements."},
        {"filename": "listP21.py", "description": "Removes duplicate sublists from a list of lists."},
        {"filename": "listP22.py", "description": "Removes consecutive duplicate elements from a list."},
        {"filename": "listP23.py", "description": "Creates an n×m matrix initialized with zeros using list comprehension."},
        {"filename": "repetedElement.py", "description": "Finds and displays repeated elements in a list."},
        {"filename": "example.py", "description": "Finds the maximum element in a list using a loop."}
    ],
    "10. Tuples": [
        {"filename": "tupleBasics.py", "description": "Demonstrates converting a list to a tuple using tuple() function."},
        {"filename": "tupleCreate.py", "description": "Shows tuple creation with different data types and conversion to list."},
        {"filename": "tupleCount.py", "description": "Counts the number of tuples present in a list."},
        {"filename": "tupleEven.py", "description": "Filters a tuple to keep only even numbers."},
        {"filename": "tupleIndex.py", "description": "Displays each element of a tuple along with its index."},
        {"filename": "tupleP1.py", "description": "Calculates the sum of elements in each tuple within a list."},
        {"filename": "tupleP2.py", "description": "Calculates the sum of all elements across lists inside a tuple."},
        {"filename": "tupleP4.py", "description": "Flattens a list of tuples into a single list."},
        {"filename": "tupleP5.py", "description": "Replaces the last element of each tuple in a list with a given number."},
        {"filename": "tupleP6.py", "description": "Removes tuples containing 'none' from a list using list comprehension."},
        {"filename": "tupleP7.py", "description": "Creates a matrix where each row is stored as a tuple."},
        {"filename": "tupleP8.py", "description": "Checks if a tuple exists in a list."},
        {"filename": "tupleP9.py", "description": "Converts all tuples in a list to lists."},
        {"filename": "tupleProduct.py", "description": "Calculates the product of all elements in a tuple."},
        {"filename": "tupleRepeated.py", "description": "Finds and displays repeated elements in a tuple."},
        {"filename": "tupleToString.py", "description": "Joins tuple elements into a single string using join()."},
        {"filename": "tupleUnpack.py", "description": "Demonstrates tuple unpacking into individual variables."}
    ],
    "11. Sets": [
        {"filename": "setsBasics.py", "description": "Demonstrates set creation using curly braces, set() constructor, and from strings."},
        {"filename": "setOperations.py", "description": "Shows set operations: union, intersection, difference, symmetric difference, subset, superset, add, remove, clear, copy."},
        {"filename": "SetREacding.py", "description": "Creates a set of tuples from user input and displays set methods using dir()."},
        {"filename": "SETPRACTICE.py", "description": "Contains multiple set practice problems: counting unique elements, finding common elements, removing duplicates, symmetric difference, and more."},
        {"filename": "setP1.py", "description": "Finds the union of two lists using set operations."}
    ],
    "12. Dictionaries": [
        {"filename": "dictionaries.py", "description": "Creates a dictionary from two lists by mapping elements as key-value pairs."},
        {"filename": "DictInbuild.py", "description": "Demonstrates dictionary methods: keys(), values(), items(), adding elements, update(), and pop()."},
        {"filename": "DictinList.py", "description": "Creates a list of dictionaries for student records with name, roll, and grade."},
        {"filename": "dictInsert.py", "description": "Creates a dictionary from user input using zip() to combine keys and values lists."},
        {"filename": "dictDelete.py", "description": "Searches for and deletes a student record from a list of dictionaries by name."},
        {"filename": "dictSearch.py", "description": "Searches for a student by name in a list of dictionaries."},
        {"filename": "dictSearchMax.py", "description": "Finds the student with the highest grade from a list of dictionaries."},
        {"filename": "dictProb1.py", "description": "Contains multiple dictionary problems: creating dict from squared values, mapping lists, word length dict, filtering by value, and even/odd classification."}
    ],
    "13. Math Operations": [
        {"filename": "mathfunc.py", "description": "Demonstrates math module functions: round, ceil, floor, abs, pow, sqrt, log, factorial, sin, cos, tan, gcd, lcm, exp, degrees, radians."},
        {"filename": "simpleintrest.py", "description": "Calculates simple interest using formula SI = (P × R × T) / 100."},
        {"filename": "compundintrest.py", "description": "Calculates compound interest and final amount using formula A = P(1 + R/100)^T."},
        {"filename": "SUMf~l.py", "description": "Calculates the sum of first and last digits of a number using while loop and string slicing."},
        {"filename": "libraries.py", "description": "Demonstrates datetime, os, random, and statistics modules with various functions."}
    ],
    "14. Matrix Operations": [
        {"filename": "addMatrices.py", "description": "Adds two matrices by taking row-wise input and computing element-wise sum."},
        {"filename": "addMAtrix.py", "description": "Adds two predefined 3×3 matrices and stores result in a list."},
        {"filename": "multiplicationMAtrices.py", "description": "Multiplies two matrices using nested loops for matrix multiplication algorithm."}
    ],
    "15. Number Problems": [
        {"filename": "Armstrong.py", "description": "Checks if a number is an Armstrong number (sum of digits^n equals the number)."},
        {"filename": "fibonacci.py", "description": "Generates Fibonacci sequence up to n terms."},
        {"filename": "neon.py", "description": "Checks if a number is a neon number (sum of digits of square equals the number)."},
        {"filename": "harshadnumber.py", "description": "Checks if a number is a Harshad number (divisible by sum of its digits)."},
        {"filename": "pallendrome.py", "description": "Checks if a number is a palindrome by reversing digits using while loop."},
        {"filename": "pallendrome(1).py", "description": "Checks if a string/number is a palindrome using string slicing [::-1]."},
        {"filename": "Check_no.py", "description": "Comprehensive function to check multiple number properties: sum of digits, neon, Harshad, Armstrong, and prime."},
        {"filename": "EvenSUm.py", "description": "Calculates the sum of even digits in a number."},
        {"filename": "numberLen.py", "description": "Counts the number of digits in a number using a while loop."},
        {"filename": "Identify.py", "description": "Recursive function attempting to identify prime numbers."},
        {"filename": "prob.py", "description": "Recursive function to find the next palindrome number from a given input."},
        {"filename": "prob1.py", "description": "Displays all prime numbers from 3 to 100."},
        {"filename": "prop2.py", "description": "Finds two indices in a list whose elements sum to a target value (Two Sum problem)."}
    ],
    "16. Temperature Conversion": [
        {"filename": "~C-F.py", "description": "Converts Fahrenheit to Celsius using formula C = (5/9) × (F - 32)."},
        {"filename": "~F-C.py", "description": "Converts Celsius to Fahrenheit using formula F = (9/5) × C + 32."}
    ]
}

# Create a new Word document
doc = Document()

# Add title
title = doc.add_heading('Python Code Collection', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add subtitle
subtitle = doc.add_paragraph('A comprehensive reference guide for Python programming')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Get the current directory
current_dir = r"c:\Users\saima\OneDrive\Documents\Sem 2\learning python"

total_files = 0

# Add Table of Contents
doc.add_heading('Table of Contents', level=1)
for category in categories.keys():
    doc.add_paragraph(category, style='List Bullet')
doc.add_page_break()

# Process each category
for category, files in categories.items():
    # Add category heading
    doc.add_heading(category, level=1)
    
    for file_info in files:
        filename = file_info["filename"]
        description = file_info["description"]
        filepath = os.path.join(current_dir, filename)
        
        # Add filename as subheading
        doc.add_heading(filename, level=2)
        
        # Add description in italic
        desc_para = doc.add_paragraph()
        desc_run = desc_para.add_run(f"Description: {description}")
        desc_run.italic = True
        desc_run.font.size = Pt(11)
        
        # Read and add file content
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if content.strip():
                # Add "Code:" label
                doc.add_paragraph("Code:", style='Intense Quote')
                
                # Add code content
                code_para = doc.add_paragraph()
                code_run = code_para.add_run(content)
                code_run.font.name = 'Consolas'
                code_run.font.size = Pt(9)
                total_files += 1
            else:
                doc.add_paragraph("(Empty file)")
        except FileNotFoundError:
            doc.add_paragraph(f"(File not found: {filename})")
        except Exception as e:
            doc.add_paragraph(f"Error reading file: {e}")
        
        # Add separator
        doc.add_paragraph("─" * 50)

# Save the document
output_path = os.path.join(current_dir, 'Python_Code_Reference.docx')
doc.save(output_path)
print(f"Document saved to: {output_path}")
print(f"Total files included: {total_files}")
